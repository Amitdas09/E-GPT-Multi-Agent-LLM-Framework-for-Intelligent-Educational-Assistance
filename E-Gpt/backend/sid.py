# backend/sid.py
import re
import io
import json
import requests
from fpdf import FPDF
from docx import Document

# ---------------------------------------------------------------------
# ⚙️ Ollama Call (4 GB-friendly)
# ---------------------------------------------------------------------
def call_ollama(prompt, model="llama3:8b", timeout=120):
    """Optimized Ollama call for stability + speed on low-RAM machines."""
    res = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.2,
                "top_p": 0.9,
                "repeat_penalty": 1.15,

                # 🔥 BIG FIX: reduce token generation
                "num_predict": 800,

                # 🔥 avoid extreme context usage
                "num_ctx": 2048,
            },
        },
        timeout=timeout,
    )
    res.raise_for_status()
    data = res.json()
    return (data.get("response") or data.get("output") or "").strip()


# ---------------------------------------------------------------------
# 🧩 Helpers
# ---------------------------------------------------------------------
def safe_json_extract(text):
    """Robust JSON extractor for LLM output with defensive repairs."""

    if not text:
        return []

    # Try direct JSON
    try:
        return json.loads(text)
    except:
        pass

    # Extract only JSON array
    match = re.search(r'\[[\s\S]*\]', text)
    if not match:
        return []

    candidate = match.group(0)

    # -------- JSON Repairs ----------
    # quote keys (A,B,C,D)
    candidate = re.sub(r'(?<!")\b([A-D])\b(?=\s*:)', r'"\1"', candidate)

    # quote field names
    candidate = re.sub(r'(?m)^\s*([a-zA-Z_]+)\s*:', r'"\1":', candidate)

    # remove trailing commas
    candidate = re.sub(r',\s*([}\]])', r'\1', candidate)

    # remove \u0000-like garbage
    candidate = candidate.replace('\x00', '')

    # final repair attempt
    try:
        return json.loads(candidate)
    except:
        try:
            from json_repair import repair_json
            repaired = repair_json(candidate)
            return json.loads(repaired)
        except:
            return []


def extract_questions_fallback(text):
    """Fallback parser when JSON parsing fails."""
    questions = []
    # Try to find question patterns
    patterns = [
        r'"question":\s*"([^"]+)"',
        r'question:\s*"([^"]+)"',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            for q_text in matches:
                questions.append({
                    "question": q_text,
                    "type": "short",
                    "marks": 1,
                    "explanation": "Answer based on course material"
                })
            break
    
    return questions if questions else []

def make_pdf(title, text):
    """Return PDF bytes."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.multi_cell(0, 10, title)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 8, "")
    for line in text.splitlines():
        safe = line.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 8, safe)
    return pdf.output(dest="S").encode("latin-1", "replace")

def make_docx(title, text):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------------------------------------------------------------
# 📝 Answer Generation
# ---------------------------------------------------------------------
def generate_answer_for_question(question_text, content, marks):
    """Generate a quality answer for a specific question."""
    mark_int = int(marks)
    
    # Determine answer length and depth
    if mark_int <= 2:
        length_guide = "2-3 clear sentences"
        depth = "concise"
    elif mark_int <= 6:
        length_guide = f"{mark_int * 2} sentences with examples"
        depth = "detailed with key concepts"
    else:
        length_guide = f"comprehensive {mark_int * 15}-word answer"
        depth = "in-depth with examples and analysis"
    
    prompt = f"""Based on this course material:

{content[:4000]}

Question: {question_text}

Provide a {depth} answer worth {marks} marks. Write {length_guide}.

Structure your answer with:
1. Direct response to the question
2. Key concepts/definitions (if applicable)
3. Examples or applications (if applicable)

Answer:"""
    
    try:
        answer = call_ollama(prompt, timeout=120)
        
        # Validate answer quality
        if answer and len(answer.strip()) > 30:
            return answer.strip()
        else:
            return f"[{marks}-mark answer] Please refer to course material on: {question_text[:50]}..."
    
    except Exception as e:
        print(f"⚠️ Answer generation error: {e}")
        return f"[{marks}-mark answer] Refer to course material for detailed explanation."

# ---------------------------------------------------------------------
# 📄 Improved Summarizer with chunking
# ---------------------------------------------------------------------
def summarize_content(content, target_length=5000):
    """Summarize long syllabus text with better chunking."""
    if len(content) < target_length:
        return content
    
    # ✅ Use extractive summarization - keep key sentences
    sentences = content.split('.')
    # Keep first and last parts, and evenly distributed middle parts
    keep_count = target_length // 100  # rough estimate of sentences
    
    if len(sentences) <= keep_count:
        return content
    
    # Keep important sections
    step = len(sentences) // keep_count
    selected = sentences[::step][:keep_count]
    
    summary = '. '.join(selected)
    
    # If still too long, use AI summarization
    if len(summary) > target_length:
        prompt = (
            "Summarize the following syllabus into key concepts and topics. "
            "Keep all important terms and concepts:\n\n" + content[:8000]
        )
        summary = call_ollama(prompt)
    
    return summary or content[:target_length]

# ---------------------------------------------------------------------
# 🧠 Build per-mark question prompt with better structure
# ---------------------------------------------------------------------
def build_mark_prompt(content, mark, count):
    """Build prompt with better structure and less truncation."""
    # ✅ Use more content (6000 chars instead of 3000)
    content_chunk = content[:6000] if len(content) > 6000 else content
    
    # ✅ Dynamic answer length guidance
    if int(mark) <= 2:
        answer_guide = "2-3 sentence explanations"
    elif int(mark) <= 6:
        answer_guide = f"{int(mark)*2} sentence explanations with key concepts"
    else:
        answer_guide = f"detailed {int(mark)*15}-word explanations with examples"
    
    return f"""You are creating exam questions. Generate EXACTLY {count} questions worth {mark} marks each.

SYLLABUS CONTENT:
{content_chunk}

REQUIREMENTS:
- Create {count} distinct {mark}-mark questions
- Questions must be clear and answerable from the content above
- Questions should test {"recall/understanding" if int(mark) <= 2 else "application/analysis" if int(mark) <= 6 else "synthesis/evaluation"}
- Each question MUST include a proper explanation/answer

OUTPUT FORMAT (CRITICAL):
<JSON_START>
[
  {{
    "question": "Clear, specific question text here",
    "type": "{"mcq" if int(mark) <= 2 else "short" if int(mark) <= 6 else "long"}",
    "marks": {mark},
    "explanation": "Provide {answer_guide} here based on the content"
  }}
]
<JSON_END>

IMPORTANT: 
- The "explanation" field is MANDATORY and must contain actual answer content, not placeholders
- Base answers on the provided syllabus content
- For {mark}-mark questions, explanations should be substantive ({answer_guide})

Generate all {count} questions with complete explanations now:"""

import difflib

def dedupe_questions_semantic(q_list):
    seen = []
    unique = []

    for q in q_list:
        qt = q.get("question", "").strip().lower()
        if not qt:
            continue

        duplicate_found = False
        for existing in seen:
            sim = difflib.SequenceMatcher(None, qt, existing).ratio()
            if sim > 0.72:
                duplicate_found = True
                break
            if qt in existing or existing in qt:
                duplicate_found = True
                break

        if not duplicate_found:
            seen.append(qt)
            unique.append(q)

    return unique



# ---------------------------------------------------------------------
# 🚀 Main Generator with batch processing
# ---------------------------------------------------------------------
def generate_paper_and_answers(syllabus_text, uploaded_files, question_counts):
    # combine content
    content = (syllabus_text or "").strip()
    if uploaded_files:
        try:
            import fitz

            for f in uploaded_files:
                f.seek(0)
                doc = fitz.open(stream=f.read(), filetype="pdf")

                for page in doc:
                    txt = page.get_text("text")
                    # skip blank pages or noisy pages
                    if txt and len(txt.strip()) > 5:
                        content += "\n\n" + txt.strip()

        except Exception as e:
            print(f"⚠️ PDF extraction error: {e}")

    if not content:
        raise ValueError("No content found from syllabus or files.")

    # ✅ Less aggressive summarization
    content = summarize_content(content, target_length=8000)

    all_questions = {}
    
    # ✅ Process questions in batches to avoid truncation
    for mark, count in question_counts.items():
        count_int = int(count)
        if count_int <= 0:
            continue
        
        # ✅ Generate in smaller batches if count is high
        batch_size = 3 if count_int > 5 else count_int
        batches = []
        
        for i in range(0, count_int, batch_size):
            batch_count = min(batch_size, count_int - i)
            prompt = build_mark_prompt(content, mark, batch_count)
            
            print(f"🔄 Generating {batch_count} x {mark}-mark questions...")
            raw = call_ollama(prompt)
            parsed = safe_json_extract(raw)
            
            if not isinstance(parsed, list):
                print(f"⚠️ {mark}-mark parse failed. Raw preview:\n{raw[:400]}")
                parsed = []
            
            # Add parsed questions and remove duplicates immediately
            batches.extend(parsed)
            batches = dedupe_questions_semantic(batches)


        
        # Final dedupe for entire mark group
        all_questions[str(mark)] = dedupe_questions_semantic(batches)

        print(f"✅ Generated {len(batches)} x {mark}-mark questions")

    # ---------- Build Question Paper ----------
    qnum = 1
    paper_lines = ["QUESTION PAPER\n", "="*60, "\n"]
    
    for mark in ["1", "2", "4", "6", "8", "10"]:
        qs = all_questions.get(mark, [])
        if not qs:
            continue
        
        paper_lines.append(f"\n{'='*60}")
        paper_lines.append(f"SECTION: {mark}-MARK QUESTIONS (Total: {len(qs)} questions)")
        paper_lines.append(f"{'='*60}\n")
        
        for q in qs:
            qt = q.get("question", "").strip()
            paper_lines.append(f"\nQ{qnum}. [{mark} marks] {qt}")
            
            if q.get("type") == "mcq" and q.get("options"):
                for k, v in q["options"].items():
                    paper_lines.append(f"    {k}. {v}")
            
            paper_lines.append("")  # blank line
            qnum += 1

    paper_text = "\n".join(paper_lines)

    # ---------- Build Answer Key ----------
    key_lines = ["ANSWER KEY\n", "="*60, "\n"]
    qnum = 1
    
    for mark in ["1", "2", "4", "6", "8", "10"]:
        qs = all_questions.get(mark, [])
        if not qs:
            continue
            
        key_lines.append(f"\n{'='*60}")
        key_lines.append(f"{mark}-MARK ANSWERS")
        key_lines.append(f"{'='*60}\n")
        
        for q in qs:
            question_text = q.get('question', '')
            ans = ""
            
            # ✅ Check if explanation already exists and is substantial
            if q.get("explanation") and len(q.get("explanation", "").strip()) > 30:
                ans = q["explanation"].strip()
                # Validate it's not a placeholder
                if "based on course material" in ans.lower() or "answer:" in ans.lower():
                    print(f"🔄 Regenerating answer for Q{qnum} (placeholder detected)...")
                    ans = generate_answer_for_question(question_text, content, mark)
            
            # ✅ For MCQs, show correct answer with explanation
            elif q.get("type") == "mcq" and q.get("correct"):
                correct_option = q.get("correct")
                correct_text = q.get("options", {}).get(correct_option, "")
                ans = f"Correct Answer: {correct_option}. {correct_text}"
                if q.get("explanation"):
                    ans += f"\n\nExplanation: {q['explanation']}"
            
            # ✅ Generate answer from content if explanation is missing/poor
            else:
                print(f"🔄 Generating answer for Q{qnum} ({mark} marks)...")
                ans = generate_answer_for_question(question_text, content, mark)
            
            key_lines.append(f"Q{qnum}. {ans}\n")
            qnum += 1

    answer_text = "\n".join(key_lines)

    # ---------- Outputs ----------
    paper_pdf = make_pdf("Generated Question Paper", paper_text)
    paper_docx = make_docx("Generated Question Paper", paper_text)
    answer_pdf = make_pdf("Answer Key", answer_text)

    return {
        "paper_text": paper_text,
        "answer_text": answer_text,
        "paper_pdf": paper_pdf,
        "paper_docx": paper_docx,
        "answer_pdf": answer_pdf
    }

# ---------------------------------------------------------------------
# 🧾 Streamlit Interface (if run standalone)
# ---------------------------------------------------------------------
if __name__ == "__main__":
    import streamlit as st
    st.title("🧾 AI Question Paper & Answer Key Generator (4 GB Optimized)")

    uploaded_files = st.file_uploader("Upload PDFs", type=["pdf"], accept_multiple_files=True)
    syllabus_text = st.text_area("Paste syllabus / content here", height=200, key="syllabus_text_area")


    st.subheader("⚙️ Marks Distribution")
    one = st.number_input("1-Mark", 0, 20, 5)
    two = st.number_input("2-Mark", 0, 20, 3)
    four = st.number_input("4-Mark", 0, 10, 3)
    six = st.number_input("6-Mark", 0, 10, 2)
    eight = st.number_input("8-Mark", 0, 10, 2)
    ten = st.number_input("10-Mark", 0, 10, 1)
    question_counts = {"1": one, "2": two, "4": four, "6": six, "8": eight, "10": ten}

    if st.button("🚀 Generate Paper & Answer Key"):
        with st.spinner("Generating..."):
            try:
                result = generate_paper_and_answers(syllabus_text, uploaded_files, question_counts)
                st.success("✅ Done!")

                st.download_button("📄 Download Paper (PDF)", result["paper_pdf"],
                                   "QuestionPaper.pdf", mime="application/pdf")
                st.download_button("📘 Download Paper (DOCX)", result["paper_docx"],
                                   "QuestionPaper.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.download_button("🧩 Download Answer Key (PDF)", result["answer_pdf"],
                                   "AnswerKey.pdf", mime="application/pdf")

                st.subheader("📘 Preview")
                with st.expander("View Question Paper Preview"):
                    st.text(result["paper_text"][:2000] + "..." if len(result["paper_text"]) > 2000 else result["paper_text"])

            except Exception as e:
                st.error(f"⚠️ {e}")
                import traceback
                st.code(traceback.format_exc())